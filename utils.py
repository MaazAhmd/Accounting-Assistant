import io
import os
from collections import defaultdict
from datetime import datetime
from io import BytesIO
from flask import send_file,make_response
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.styles import getSampleStyleSheet
from docx.enum.table import WD_TABLE_ALIGNMENT
from models import Transaction, db,Category
from flask_login import current_user  
import pandas as pd
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle,Paragraph
import tempfile
from difflib import SequenceMatcher
import re






def seed_categories():
    asset_categories = [
        "Intangible Assets", "Fixed Assets", "Long-Term Financial Assets",
        "Deferred Taxes", "Inventory", "Receivables", "Receivables Over One Year",
        "Investments", "Cash", "Prepaid Expenses", "Unpaid Capital"
    ]

    liability_categories = [
        "Issued Capital", "Share Premiums", "Revaluation Reserve", "Reserves",
        "Retained Earnings", "Current Profit/Loss", "Provisions",
        "Liabilities Under One Year", "Liabilities Over One Year", "Deferred Income",
        "Liabilities Debit", "Liabilities Credit"
    ]

    # Add asset categories
    for category_name in asset_categories:
        if not Category.query.filter_by(name=category_name, type="income").first():
            db.session.add(Category(name=category_name, type="income"))

    # Add liability categories
    for category_name in liability_categories:
        if not Category.query.filter_by(name=category_name, type="expense").first():
            db.session.add(Category(name=category_name, type="expense"))

    db.session.commit()
    print("Categories seeded successfully.")


def автоматично_дефинирана_категория(transaction_description):
    # Проста логика за определяне на категорията на база описание
    if "заплата, salary" in transaction_description.lower() or "персонал, staff" in transaction_description.lower():
        return 'Personnel expenses'
    elif "материали, materials, supplies" in transaction_description.lower() or "доставки, supplies" in transaction_description.lower():
        return 'Raw materials, supplies, and external services expenses'
    elif "амортизация, depreciation" in transaction_description.lower():
        return 'Depreciation and amortization expenses'
    elif "данък, tax" in transaction_description.lower():
        return 'Tax expenses'
    elif "приход,revenue" in transaction_description.lower() or "продажба, sales, sale" in transaction_description.lower():
        return 'Net sales revenue'
    elif "друг приход, other income" in transaction_description.lower():
        return 'Other revenue'  # Добавена логика за другите приходи
    else:
        return 'Other expenses'  # Ако няма съвпадение, по подразбиране ще бъде 'Other expenses'

def recalculate_totals():
    # Извличане на всички транзакции на текущия потребител
    transactions = Transaction.query.filter_by(user_id=current_user.id).all()

    # Пресмятане на общите суми за дебит и кредит
    asset_debit_total = 0.0
    asset_credit_total = 0.0
    liabilities_debit_total = 0.0
    liabilities_credit_total = 0.0

    for transaction in transactions:
        if not transaction.credit:
            if transaction.type =='income':
                asset_debit_total += transaction.amount
            elif transaction.type=='expense':
                liabilities_debit_total += transaction.amount

        if transaction.credit:
            if transaction.type == 'income':
                asset_credit_total += transaction.amount
            elif transaction.type =='expense':
                liabilities_credit_total += transaction.amount

    # Записване на стойностите в базата данни
    current_user.asset_debit_total = asset_debit_total
    current_user.asset_credit_total = asset_credit_total
    current_user.liabilities_debit_total = liabilities_debit_total
    current_user.liabilities_credit_total = liabilities_credit_total

    db.session.commit()
    
def normalize_text(text):
    """Normalize text by removing extra spaces, commas, and converting to lowercase."""
    return re.sub(r'[,\s]+', ' ', text.strip()).lower()

def match_category(category, keyword_mapping):
    """Match a category to the keyword mapping using normalized text."""
    normalized_category = normalize_text(category)
    for keywords, mapped_value in keyword_mapping.items():
        if any(keyword in normalized_category for keyword in keywords):
            return mapped_value
    return None

def calculate_income_expense_data(transactions):
    # Initialize the income_expense_data dictionary
    income_expense_data = defaultdict(float)

    # Define keyword mapping
    keyword_mapping = {
        (normalize_text("raw materials"), normalize_text("supplies"), normalize_text("external services expenses")): "_3_raw_material_expenses",
        (normalize_text("personnel expenses"),): "_4_personnel_expenses",
        (normalize_text("depreciation expenses"),): "_5_depreciation_expenses",
        (normalize_text("other expenses"),): "_6_other_expenses",
        (normalize_text("tax expenses"),): "_7_tax_expenses",
        (normalize_text("net sales revenue"),): "_1_net_sales_revenue",
        (normalize_text("other revenue"),): "_2_other_revenue",
    }

    # Get the current and previous year
    current_year = datetime.now().year
    previous_year = current_year - 1

    # Group transactions by category and year
    for transaction in transactions:
        year = transaction.date.year
        category = transaction.income_statement_category or ""
        amount = transaction.amount

        # Normalize and match category
        mapped_category = match_category(category, keyword_mapping)
        
        if year == current_year and mapped_category:
            income_expense_data[f"{mapped_category}_current"] += amount
        elif year == previous_year and mapped_category:
            income_expense_data[f"{mapped_category}_previous"] += amount

    # Calculate totals and derived metrics
    income_expense_data["_total_expenses_current"] = (
        income_expense_data["_3_raw_material_expenses_current"] +
        income_expense_data["_4_personnel_expenses_current"] +
        income_expense_data["_5_depreciation_expenses_current"] +
        income_expense_data["_6_other_expenses_current"]
    )
    income_expense_data["_total_expenses_previous"] = (
        income_expense_data["_3_raw_material_expenses_previous"] +
        income_expense_data["_4_personnel_expenses_previous"] +
        income_expense_data["_5_depreciation_expenses_previous"] +
        income_expense_data["_6_other_expenses_previous"]
    )
    income_expense_data["_total_revenue_current"] = (
        income_expense_data["_1_net_sales_revenue_current"] +
        income_expense_data["_2_other_revenue_current"]
    )
    income_expense_data["_total_revenue_previous"] = (
        income_expense_data["_1_net_sales_revenue_previous"] +
        income_expense_data["_2_other_revenue_previous"]
    )

    income_expense_data["_9_accounting_loss_current"] = max(
        0,
        income_expense_data["_total_expenses_current"] - income_expense_data["_total_revenue_current"]
    )
    income_expense_data["_9_accounting_loss_previous"] = max(
        0,
        income_expense_data["_total_expenses_previous"] - income_expense_data["_total_revenue_previous"]
    )


    income_expense_data["_11_total_loss_current"] = max(
        0,
        income_expense_data["_total_expenses_current"] +income_expense_data["_7_tax_expenses_current"]- income_expense_data["_total_revenue_current"]
    )
    income_expense_data["_11_total_loss_previous"] = max(
        0,
        income_expense_data["_total_expenses_previous"] +income_expense_data["_7_tax_expenses_current"] - income_expense_data["_total_revenue_previous"]
    )



    income_expense_data["_8_accounting_profit_current"] = max(
        0,
        income_expense_data["_total_revenue_current"] - income_expense_data["_total_expenses_current"]
    )
    income_expense_data["_8_accounting_profit_previous"] = max(
        0,
        income_expense_data["_total_revenue_previous"] - income_expense_data["_total_expenses_previous"]
    )



    income_expense_data["_total_all_expenses_current"] = (
        income_expense_data["_total_expenses_current"] +
        income_expense_data["_7_tax_expenses_current"] +
        income_expense_data["_10_net_profit_current"]
    )
    income_expense_data["_total_all_expenses_previous"] = (
        income_expense_data["_total_expenses_previous"] +
        income_expense_data["_7_tax_expenses_previous"] +
        income_expense_data["_10_net_profit_previous"]
    )

    income_expense_data["_total_all_revenue_current"] = (
        income_expense_data["_total_revenue_current"] +
        income_expense_data.get("_11_total_loss_current", 0)
    )
    income_expense_data["_total_all_revenue_previous"] = (
        income_expense_data["_total_revenue_previous"] +
        income_expense_data.get("_11_total_loss_previous", 0)
    )

    income_expense_data["_10_net_profit_current"] = max(0,
        income_expense_data.get("_total_all_revenue_current", 0) -
        income_expense_data["_total_all_expenses_current"]-income_expense_data["_7_tax_expenses_current"]
    )
    income_expense_data["_10_net_profit_previous"] = max(0,
        income_expense_data.get("_total_all_revenue_previous", 0) -
        income_expense_data["_total_all_expenses_previous"]-income_expense_data["_7_tax_expenses_current"]
    )
    return income_expense_data

def calculate_asset_data(transactions):
    from collections import defaultdict
    
    # Initialize the asset_data dictionary
    asset_data = defaultdict(float)

    # Get the current and previous year
    current_year = datetime.now().year
    previous_year = current_year - 1

    # category_mapping = {
    #     "subscribed but unpaid capital": "A_asset_unpaid_capital",
    #     "intangible assets": "B_intangible_assets",
    #     "property, plant, and equipment": "B_fixed_assets",
    #     "long-term financial assets": "B_long_term_financial_assets",
    #     "deferred taxes": "B_deferred_taxes",
    #     "inventory": "C_inventory",
    #     "receivables": "C_receivables",
    #     "receivables over one year": "C_receivables_over_one_year",
    #     "investments": "C_investments",
    #     "cash": "C_cash",
    #     "prepaid expenses": "D_prepaid_expenses",
    # }

    # Define keyword mapping for categories
    keyword_mapping = [
        ("A_asset_unpaid_capital", ["subscribed", "unpaid", "capital"]),
        ("B_intangible_assets", ["intangible", "assets"]),
        ("B_fixed_assets", ["property", "plant", "equipment"]),
        ("B_long_term_financial_assets", ["long-term", "financial", "assets"]),
        ("B_deferred_taxes", ["deferred", "taxes"]),
        ("B_deferred_taxes", ["defered", "taxes"]),
        ("C_inventory", ["inventory"]),
        ("C_receivables", ["receivable"]),
        ("C_receivables_over_one_year", ["receivable", "over", "one", "year"]),
        ("C_receivables_over_one_year", ["receivable", "over", "1", "year"]),
        ("C_investments", ["investments"]),
        ("C_cash", ["cash"]),
        ("D_prepaid_expenses", ["prepaid", "expenses"]),
    ]

    def similar(a, b, threshold=0.8):
        return SequenceMatcher(None, a, b).ratio() > threshold

    def map_category_to_key(category):
        category_lower = category.lower()
        matched_keys = []

        for key, keywords in keyword_mapping:
            if all(any(similar(re.sub(r'\W+', '', word), keyword) for word in category_lower.split()) for keyword in keywords):
                matched_keys.append(key)

        if matched_keys:
            return matched_keys[0]
        return None

    for transaction in transactions:
        year = transaction.date.year
        debit_category = transaction.debit.lower() if transaction.debit else ""
        credit_category = transaction.credit.lower() if transaction.credit else ""
        amount = transaction.amount

        if year == current_year:
            suffix = "_current"
        elif year == previous_year:
            suffix = "_previous"
        else:
            continue  
        
        if "provisions" in credit_category and not debit_category:
            asset_data['D_prepaid_expenses' + suffix] += amount
        
        else:
            debit_key = map_category_to_key(debit_category)
            credit_key = map_category_to_key(credit_category)

            if debit_key:
                asset_data[debit_key + suffix] += amount
            if credit_key:
                # if credit_category=="cash":
                #      asset_data[credit_key+ suffix] -= amount
                # else:
                asset_data[credit_key + suffix] -= amount

    asset_data['B_total_noncurrent_assets_current'] = (
        asset_data['B_intangible_assets_current'] +
        asset_data['B_fixed_assets_current'] +
        asset_data['B_long_term_financial_assets_current'] +
        asset_data['B_deferred_taxes_current']
    )
    asset_data['B_total_noncurrent_assets_previous'] = (
        asset_data['B_intangible_assets_previous'] +
        asset_data['B_fixed_assets_previous'] +
        asset_data['B_long_term_financial_assets_previous'] +
        asset_data['B_deferred_taxes_previous']
    )
    asset_data['C_total_current_assets_current'] = (
        asset_data['C_inventory_current'] +
        asset_data['C_receivables_current'] +
        asset_data['C_receivables_over_one_year_current'] +
        asset_data['C_investments_current'] +
        asset_data['C_cash_current']
    )
    asset_data['C_total_current_assets_previous'] = (
        asset_data['C_inventory_previous'] +
        asset_data['C_receivables_previous'] +
        asset_data['C_receivables_over_one_year_previous'] +
        asset_data['C_investments_previous'] +
        asset_data['C_cash_previous']
    )
    asset_data['total_assets_current'] = (
        asset_data['A_asset_unpaid_capital_current'] +
        asset_data['B_total_noncurrent_assets_current'] +
        asset_data['C_total_current_assets_current'] +
        asset_data['D_prepaid_expenses_current']
    )
    asset_data['total_assets_previous'] = (
        asset_data['A_asset_unpaid_capital_previous'] +
        asset_data['B_total_noncurrent_assets_previous'] +
        asset_data['C_total_current_assets_previous'] +
        asset_data['D_prepaid_expenses_previous']
    )

    return dict(asset_data)


def calculate_liability_data(transactions):
    from collections import defaultdict
    
    # Initialize the liability_data dictionary
    liability_data = defaultdict(float)

    # Get the current and previous year
    current_year = datetime.now().year
    previous_year = current_year - 1

    # category_mapping = {
    #     "issued capital": "A_issued_capital",
    #     "share premiums": "A_share_premiums",
    #     "revaluation reserve": "A_revaluation_reserve",
    #     "reserves": "A_reserves",
    #     "retained earnings": "A_retained_earnings",
    #     "current profit/loss": "A_current_profit_loss",
    #     "shareholder's equity": "A_shareholders_equity",
    #     "shareholder equity": "A_shareholders_equity",
    #     "shareholders equity": "A_shareholders_equity",
    #     "shareholders' equity": "A_shareholders_equity",
    #     "shareholders’ equity": "A_shareholders_equity",
    #     "provisions": "B_provisions",
    #     "provisions and similar obligations": "B_provisions",
    #     "liabilities under one year": "C_liabilities_one_year",
    #     "liabilities up to one year": "C_liabilities_one_year",
    #     "liabilities up to 1 year": "C_liabilities_one_year",
    #     "liabilities under 1 year": "C_liabilities_one_year",
    #     "liabilities over one year": "C_liabilities_over_one_year",
    #     "liabilities over 1 year": "C_liabilities_over_one_year",
    #     "deferred taxes": "E_deferred_tax_liabilities",
    #     "deferred tax": "E_deferred_tax_liabilities",
    #     "deferred income": "D_deferred_income",
    #     "deffered income": "D_deferred_income",
    #     "liabilities debit": "liabilities_debit",
    #     "liabilities credit": "liabilities_credit",
    # }

    # Define keyword mapping for liabilities
    keyword_mapping = [
        ("A_issued_capital", ["issued", "capital"]),
        ("A_share_premiums", ["share", "premium"]),
        ("A_revaluation_reserve", ["revaluation", "reserve"]),
        ("A_reserves", ["reserves"]),
        ("A_retained_earnings", ["retained", "earnings"]),
        ("A_current_profit_loss", ["current", "profit"]),
        ("A_current_profit_loss", ["revenue"]),
        ("A_shareholders_equity", ["shareholder", "equity"]),
        ("B_provisions", ["provisions"]),
        ("C_liabilities_one_year", ["liabilities", "under", "one", "year"]),
        ("C_liabilities_one_year", ["liabilities", "under", "1", "year"]),
        ("C_liabilities_one_year", ["liabilities", "up", "one", "year"]),
        ("C_liabilities_one_year", ["liabilities", "up", "1", "year"]),
        ("C_liabilities_over_one_year", ["liabilities", "over", "one", "year"]),
        ("C_liabilities_over_one_year", ["liabilities", "over", "1", "year"]),
        ("E_deferred_tax_liabilities", ["deferred", "tax", "liabilities"]),
        ("D_deferred_income", ["deferred", "income"]),
        ("D_deferred_income", ["deffered", "income"]),
        ("D_deferred_income", ["defferred", "income"]),
        ("D_deferred_income", ["defered", "income"]),
        ("liabilities_debit", ["liabilities", "debit"]),
        ("liabilities_credit", ["liabilities", "credit"]),
    ]

    def similar(a, b, threshold=0.8):
        return SequenceMatcher(None, a, b).ratio() > threshold

    def map_category_to_key(category):
        category_lower = category.lower()
        matched_keys = []

        for key, keywords in keyword_mapping:
            if all(any(similar(re.sub(r'\W+', '', word), keyword) for word in category_lower.split()) for keyword in keywords):
                # print("Key is ", key)
                # print("Keywords are ", keywords)
                # print("Word is ", category_lower)
                # print("Matched")
                matched_keys.append(key)
            # else:
            #     print("Not matched: ")
            #     print("Key is ", key)
            #     print("Keywords are ", keywords)
            #     print("Word is ", category_lower)

        if matched_keys:
            return matched_keys[0]
        return None

    for transaction in transactions:
        year = transaction.date.year
        debit_category = transaction.debit.lower() if transaction.debit else ""
        credit_category = transaction.credit.lower() if transaction.credit else ""
        amount = transaction.amount

        if year == current_year:
            suffix = "_current"
        elif year == previous_year:
            suffix = "_previous"
        else:
            continue  

        if "deferred" in debit_category and "tax" in debit_category and credit_category == "":
            liability_data['E_deferred_tax_liabilities' + suffix] += amount
        elif "pre" in debit_category and "expense" in debit_category  and "paid" in debit_category and not credit_category:
            liability_data['C_liabilities_one_year' + suffix] += amount
        elif ("deffered" in debit_category or "deferred" in debit_category) and "income" in debit_category and not credit_category:
            liability_data['D_deferred_income' + suffix] -= amount
            liability_data['E_deferred_tax_liabilities' + suffix] += amount
        elif "receivable" in debit_category and not credit_category:
            liability_data['A_current_profit_loss' + suffix] += amount
        else:
            debit_key = map_category_to_key(debit_category)
            credit_key = map_category_to_key(credit_category)

            if debit_key:
                liability_data[debit_key + suffix] -= amount
            if credit_key:
                # if credit_category=="cash":
                #  liability_data[credit_key+ suffix] -= amount
                # else:
                liability_data[credit_key + suffix] += amount

    liability_data['A_total_equity_current'] = (
        liability_data['A_issued_capital_current'] +
        liability_data['A_share_premiums_current'] +
        liability_data['A_revaluation_reserve_current'] +
        liability_data['A_reserves_current'] +
        liability_data['A_retained_earnings_current'] +
        liability_data['A_current_profit_loss_current'] +
        liability_data['A_shareholders_equity_current']
    )
    liability_data['A_total_equity_previous'] = (
        liability_data['A_issued_capital_previous'] +
        liability_data['A_share_premiums_previous'] +
        liability_data['A_revaluation_reserve_previous'] +
        liability_data['A_reserves_previous'] +
        liability_data['A_retained_earnings_previous'] +
        liability_data['A_current_profit_loss_previous'] +
        liability_data['A_shareholders_equity_previous']
    )
    liability_data['total_liabilities_current'] = (
        liability_data['A_total_equity_current'] +
        liability_data['B_provisions_current'] +
        liability_data['C_liabilities_one_year_current'] +
        liability_data['C_liabilities_over_one_year_current'] +
        liability_data['D_deferred_income_current'] +
        liability_data['E_deferred_tax_liabilities_current']
    )
    liability_data['total_liabilities_previous'] = (
        liability_data['A_total_equity_previous'] +
        liability_data['B_provisions_previous'] +
        liability_data['C_liabilities_one_year_previous'] +
        liability_data['C_liabilities_over_one_year_previous'] +
        liability_data['D_deferred_income_previous'] +
        liability_data['E_deferred_tax_liabilities_previous']
    )

    return dict(liability_data)



def export_income_expense_pdf():
    transactions = Transaction.query.filter_by(user_id=current_user.id).all()
    income_expense_data = calculate_income_expense_data(transactions)

    # Prepare data for the table
    table_data = [
        ["Sections, Groups, Items", "Amount (thousand EUR) Current Year", "Amount (thousand EUR) Previous Year"]
    ]

    # Populate table rows
    table_data.extend([
        ["1. Net Sales Revenue", income_expense_data.get("_1_net_sales_revenue_current", 0), income_expense_data.get("_1_net_sales_revenue_previous", 0)],
        ["2. Other Revenue", income_expense_data.get("_2_other_revenue_current", 0), income_expense_data.get("_2_other_revenue_previous", 0)],
        ["Total Revenue", income_expense_data.get("_total_revenue_current", 0), income_expense_data.get("_total_revenue_previous", 0)],
        ["3. Raw Material Expenses", income_expense_data.get("_3_raw_material_expenses_current", 0), income_expense_data.get("_3_raw_material_expenses_previous", 0)],
        ["4. Personnel Expenses", income_expense_data.get("_4_personnel_expenses_current", 0), income_expense_data.get("_4_personnel_expenses_previous", 0)],
        ["5. Depreciation Expenses", income_expense_data.get("_5_depreciation_expenses_current", 0), income_expense_data.get("_5_depreciation_expenses_previous", 0)],
        ["6. Other Expenses", income_expense_data.get("_6_other_expenses_current", 0), income_expense_data.get("_6_other_expenses_previous", 0)],
        ["Total Expenses", income_expense_data.get("_total_expenses_current", 0), income_expense_data.get("_total_expenses_previous", 0)],
        ["8. Accounting Profit", income_expense_data.get("_8_accounting_profit_current", 0), income_expense_data.get("_8_accounting_profit_previous", 0)],
        ["9. Accounting Loss", income_expense_data.get("_9_accounting_loss_current", 0), income_expense_data.get("_9_accounting_loss_previous", 0)],
        ["7. Tax Expenses", income_expense_data.get("_7_tax_expenses_current", 0), income_expense_data.get("_7_tax_expenses_previous", 0)],
        ["10. Net Profit", income_expense_data.get("_10_net_profit_current", 0), income_expense_data.get("_10_net_profit_previous", 0)],
        ["11. Total Loss", income_expense_data.get("_11_total_loss_current", 0), income_expense_data.get("_11_total_loss_previous", 0)],
        ["Total Revenue (Including Loss)", income_expense_data.get("_total_all_revenue_current", 0), income_expense_data.get("_total_all_revenue_previous", 0)],
        ["Total Expenses (Including Taxes and Net Profit)", income_expense_data.get("_total_all_expenses_current", 0), income_expense_data.get("_total_all_expenses_previous", 0)]
    ])

    # Create PDF in memory
    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []

    # Create the table
    table = Table(table_data)
    table_style = TableStyle([
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.blue),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ])
    table.setStyle(table_style)
    elements.append(table)

    # Build the PDF
    pdf.build(elements)

    # Return the PDF response
    buffer.seek(0)
    response = make_response(buffer.getvalue())
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = 'attachment; filename=income_expense_statement.pdf'
    return response




def export_income_expense_word():
    transactions = Transaction.query.filter_by(user_id=current_user.id).all()
    income_expense_data = calculate_income_expense_data(transactions)

    # Create a Word document
    document = Document()
    document.add_heading("Income and Expense Statement", level=1)

    # Create the table
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style for the header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Sections, Groups, Items"
    header_cells[1].text = "Amount (thousand EUR) Current Year"
    header_cells[2].text = "Amount (thousand EUR) Previous Year"

    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

        # Add a subtle blue background to header cells
        tcPr = cell._element.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            cell._element.insert(0, tcPr)

        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:fill'), "ADD8E6")  # Light blue
        tcPr.append(shading_elm)

    # Populate the table rows with income/expense data
    rows = [
        ["1. Net Sales Revenue", income_expense_data.get("_1_net_sales_revenue_current", 0), income_expense_data.get("_1_net_sales_revenue_previous", 0)],
        ["2. Other Revenue", income_expense_data.get("_2_other_revenue_current", 0), income_expense_data.get("_2_other_revenue_previous", 0)],
        ["Total Revenue", income_expense_data.get("_total_revenue_current", 0), income_expense_data.get("_total_revenue_previous", 0)],
        ["3. Raw Material Expenses", income_expense_data.get("_3_raw_material_expenses_current", 0), income_expense_data.get("_3_raw_material_expenses_previous", 0)],
        ["4. Personnel Expenses", income_expense_data.get("_4_personnel_expenses_current", 0), income_expense_data.get("_4_personnel_expenses_previous", 0)],
        ["5. Depreciation Expenses", income_expense_data.get("_5_depreciation_expenses_current", 0), income_expense_data.get("_5_depreciation_expenses_previous", 0)],
        ["6. Other Expenses", income_expense_data.get("_6_other_expenses_current", 0), income_expense_data.get("_6_other_expenses_previous", 0)],
        ["Total Expenses", income_expense_data.get("_total_expenses_current", 0), income_expense_data.get("_total_expenses_previous", 0)],
        ["8. Accounting Profit", income_expense_data.get("_8_accounting_profit_current", 0), income_expense_data.get("_8_accounting_profit_previous", 0)],
        ["9. Accounting Loss", income_expense_data.get("_9_accounting_loss_current", 0), income_expense_data.get("_9_accounting_loss_previous", 0)],
        ["7. Tax Expenses", income_expense_data.get("_7_tax_expenses_current", 0), income_expense_data.get("_7_tax_expenses_previous", 0)],
        ["10. Net Profit", income_expense_data.get("_10_net_profit_current", 0), income_expense_data.get("_10_net_profit_previous", 0)],
        ["11. Total Loss", income_expense_data.get("_11_total_loss_current", 0), income_expense_data.get("_11_total_loss_previous", 0)],
        ["Total Revenue (Including Loss)", income_expense_data.get("_total_all_revenue_current", 0), income_expense_data.get("_total_all_revenue_previous", 0)],
        ["Total Expenses (Including Taxes and Net Profit)", income_expense_data.get("_total_all_expenses_current", 0), income_expense_data.get("_total_all_expenses_previous", 0)]
    ]

    for row in rows:
        cells = table.add_row().cells
        for i, data in enumerate(row):
            cells[i].text = str(data)
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)  # Set font size
            cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

    # Save the document to a temporary file
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, "income_expense_statement.docx")
    document.save(file_path)

    # Return the file as a response
    return send_file(file_path, as_attachment=True, download_name="income_expense_statement.docx")





def export_income_expense_excel():
    transactions = Transaction.query.filter_by(user_id=current_user.id).all()
    income_expense_data = calculate_income_expense_data(transactions)

    # Create a DataFrame with columns matching the HTML layout
    columns = ["Sections, Groups, Items", "Amount (thousand EUR) Current Year", "Amount (thousand EUR) Previous Year"]
    rows = [
        ["1. Net Sales Revenue", income_expense_data.get("_1_net_sales_revenue_current", 0), income_expense_data.get("_1_net_sales_revenue_previous", 0)],
        ["2. Other Revenue", income_expense_data.get("_2_other_revenue_current", 0), income_expense_data.get("_2_other_revenue_previous", 0)],
        ["Total Revenue", income_expense_data.get("_total_revenue_current", 0), income_expense_data.get("_total_revenue_previous", 0)],
        ["3. Raw Material Expenses", income_expense_data.get("_3_raw_material_expenses_current", 0), income_expense_data.get("_3_raw_material_expenses_previous", 0)],
        ["4. Personnel Expenses", income_expense_data.get("_4_personnel_expenses_current", 0), income_expense_data.get("_4_personnel_expenses_previous", 0)],
        ["5. Depreciation Expenses", income_expense_data.get("_5_depreciation_expenses_current", 0), income_expense_data.get("_5_depreciation_expenses_previous", 0)],
        ["6. Other Expenses", income_expense_data.get("_6_other_expenses_current", 0), income_expense_data.get("_6_other_expenses_previous", 0)],
        ["Total Expenses", income_expense_data.get("_total_expenses_current", 0), income_expense_data.get("_total_expenses_previous", 0)],
        ["8. Accounting Profit", income_expense_data.get("_8_accounting_profit_current", 0), income_expense_data.get("_8_accounting_profit_previous", 0)],
        ["9. Accounting Loss", income_expense_data.get("_9_accounting_loss_current", 0), income_expense_data.get("_9_accounting_loss_previous", 0)],
        ["7. Tax Expenses", income_expense_data.get("_7_tax_expenses_current", 0), income_expense_data.get("_7_tax_expenses_previous", 0)],
        ["10. Net Profit", income_expense_data.get("_10_net_profit_current", 0), income_expense_data.get("_10_net_profit_previous", 0)],
        ["11. Total Loss", income_expense_data.get("_11_total_loss_current", 0), income_expense_data.get("_11_total_loss_previous", 0)],
        ["Total Revenue (Including Loss)", income_expense_data.get("_total_all_revenue_current", 0), income_expense_data.get("_total_all_revenue_previous", 0)],
        ["Total Expenses (Including Taxes and Net Profit)", income_expense_data.get("_total_all_expenses_current", 0), income_expense_data.get("_total_all_expenses_previous", 0)]
    ]

    # Create the DataFrame
    df = pd.DataFrame(rows, columns=columns)

    # Save the DataFrame to an Excel file
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, "income_expense_statement.xlsx")
    df.to_excel(file_path, index=False)

    # Return the file as a response
    return send_file(file_path, as_attachment=True, download_name="income_expense_statement.xlsx")

def format_negative(value):
    """Format negative values with parentheses."""
    if isinstance(value, (int, float)) and value < 0:
        return f"({abs(value):,.2f})"
    return f"{value:,.2f}" if isinstance(value, (int, float)) else value

def export_balance_sheet_pdf():
    # Fetch data using the calculation functions
    transactions = Transaction.query.filter_by(user_id=current_user.id).all()

    assets = calculate_asset_data(transactions)
    liabilities = calculate_liability_data(transactions)

    # Create a PDF document
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []

    # Styles
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    normal_style = styles["Normal"]

    # Add Title
    elements.append(Paragraph("Balance Sheet", title_style))

    # **Assets Section**
    elements.append(Paragraph("Assets", styles["Heading2"]))
    asset_data = [
        ["Sections, Groups, Items", "Amount (Current Year)", "Amount (Previous Year)"]
    ]
    asset_data.extend([
        ["A. Subscribed but Unpaid Capital", format_negative(assets["A_asset_unpaid_capital_current"]), format_negative(assets["A_asset_unpaid_capital_previous"])],
        ["B. Non-Current (Long-Term) Assets", "", ""],
        ["  I. Intangible Assets", format_negative(assets["B_intangible_assets_current"]), format_negative(assets["B_intangible_assets_previous"])],
        ["  II. Property, Plant, and Equipment", format_negative(assets["B_fixed_assets_current"]), format_negative(assets["B_fixed_assets_previous"])],
        ["  III. Long-Term Financial Assets", format_negative(assets["B_long_term_financial_assets_current"]), format_negative(assets["B_long_term_financial_assets_previous"])],
        ["  IV. Deferred Taxes", format_negative(assets["B_deferred_taxes_current"]), format_negative(assets["B_deferred_taxes_previous"])],
        ["  Total for Section B", format_negative(assets["B_total_noncurrent_assets_current"]), format_negative(assets["B_total_noncurrent_assets_previous"])],
        ["C. Current (Short-Term) Assets", "", ""],
        ["  I. Inventory", format_negative(assets["C_inventory_current"]), format_negative(assets["C_inventory_previous"])],
        ["  II. Receivables", format_negative(assets["C_receivables_current"]), format_negative(assets["C_receivables_previous"])],
        ["  III. Investments", format_negative(assets["C_investments_current"]), format_negative(assets["C_investments_previous"])],
        ["  IV. Cash and Cash Equivalents", format_negative(assets["C_cash_current"]), format_negative(assets["C_cash_previous"])],
        ["  Total for Section C", format_negative(assets["C_total_current_assets_current"]), format_negative(assets["C_total_current_assets_previous"])],
        ["D. Prepaid Expenses", format_negative(assets["D_prepaid_expenses_current"]), format_negative(assets["D_prepaid_expenses_previous"])],
        ["TOTAL ASSETS", format_negative(assets["total_assets_current"]), format_negative(assets["total_assets_previous"])]
    ])
    asset_table = Table(asset_data, hAlign="LEFT")
    asset_table.setStyle(TableStyle([
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.blue),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(asset_table)

    # **Liabilities Section**
    elements.append(Paragraph("Liabilities", styles["Heading2"]))
    liability_data = [
        ["Sections, Groups, Items", "Amount (Current Year)", "Amount (Previous Year)"]
    ]
    liability_data.extend([
        ["A. Equity", "", ""],
        ["  I. Issued Capital", format_negative(liabilities["A_issued_capital_current"]), format_negative(liabilities["A_issued_capital_previous"])],
        ["  II. Share Premiums", format_negative(liabilities["A_share_premiums_current"]), format_negative(liabilities["A_share_premiums_previous"])],
        ["  III. Revaluation Reserve", format_negative(liabilities["A_revaluation_reserve_current"]), format_negative(liabilities["A_revaluation_reserve_previous"])],
        ["  IV. Reserves", format_negative(liabilities["A_reserves_current"]), format_negative(liabilities["A_reserves_previous"])],
        ["  V. Retained Earnings (Loss)", format_negative(liabilities["A_retained_earnings_current"]), format_negative(liabilities["A_retained_earnings_previous"])],
        ["  VI. Current Profit (Loss)", format_negative(liabilities["A_current_profit_loss_current"]), format_negative(liabilities["A_current_profit_loss_previous"])],
        ["  VII. Shareholders Equity", format_negative(liabilities["A_shareholders_equity_current"]), format_negative(liabilities["A_shareholders_equity_previous"])],
        ["  Total for Section A", format_negative(liabilities["A_total_equity_current"]), format_negative(liabilities["A_total_equity_previous"])],
        ["B. Provisions and Similar Obligations", format_negative(liabilities["B_provisions_current"]), format_negative(liabilities["B_provisions_previous"])],
        ["C. Liabilities", "", ""],
        ["  Up to 1 Year", format_negative(liabilities["C_liabilities_one_year_current"]), format_negative(liabilities["C_liabilities_one_year_previous"])],
        ["  Over 1 Year", format_negative(liabilities["C_liabilities_over_one_year_current"]), format_negative(liabilities["C_liabilities_over_one_year_previous"])],
        ["D. Financing and Deferred Income", format_negative(liabilities["D_deferred_income_current"]), format_negative(liabilities["D_deferred_income_previous"])],
        ["E. Deferred Tax Liabilities", format_negative(liabilities["E_deferred_tax_liabilities_current"]), format_negative(liabilities["E_deferred_tax_liabilities_previous"])],
        ["TOTAL LIABILITIES", format_negative(liabilities["total_liabilities_current"]), format_negative(liabilities["total_liabilities_previous"])]
    ])
    liability_table = Table(liability_data, hAlign="LEFT")

    liability_table.setStyle(TableStyle([
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.blue),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(liability_table)

    # Build PDF
    doc.build(elements)
    buffer.seek(0)

    # Return the response
    response = make_response(buffer.read())
    response.headers["Content-Disposition"] = "attachment; filename=balance_sheet.pdf"
    response.headers["Content-Type"] = "application/pdf"
    return response



from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_balance_sheet_word():
    # Helper function to format numbers
    def format_number(value):
        if isinstance(value, (int, float)):
            return f"({abs(value):,.2f})" if value < 0 else f"{value:,.2f}"
        return value

    # Fetch data using the calculation functions
    transactions = Transaction.query.filter_by(user_id=current_user.id).all()

    assets = calculate_asset_data(transactions)
    liabilities = calculate_liability_data(transactions)

    # Create a Word document
    doc = Document()

    # Title
    title = doc.add_heading("Balance Sheet", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # **Assets Section**
    doc.add_heading("Assets", level=2)
    asset_table = doc.add_table(rows=1, cols=3)
    asset_table.style = 'Table Grid'

    # Add headers
    headers = asset_table.rows[0].cells
    headers[0].text = "Sections, Groups, Items"
    headers[1].text = "Amount (Current Year)"
    headers[2].text = "Amount (Previous Year)"

    # Add asset data
    asset_data = [
        ["A. Subscribed but Unpaid Capital", assets["A_asset_unpaid_capital_current"], assets["A_asset_unpaid_capital_previous"]],
        ["B. Non-Current (Long-Term) Assets", "", ""],
        ["  I. Intangible Assets", assets["B_intangible_assets_current"], assets["B_intangible_assets_previous"]],
        ["  II. Property, Plant, and Equipment", assets["B_fixed_assets_current"], assets["B_fixed_assets_previous"]],
        ["  III. Long-Term Financial Assets", assets["B_long_term_financial_assets_current"], assets["B_long_term_financial_assets_previous"]],
        ["  IV. Deferred Taxes", assets["B_deferred_taxes_current"], assets["B_deferred_taxes_previous"]],
        ["  Total for Section B", assets["B_total_noncurrent_assets_current"], assets["B_total_noncurrent_assets_previous"]],
        ["C. Current (Short-Term) Assets", "", ""],
        ["  I. Inventory", assets["C_inventory_current"], assets["C_inventory_previous"]],
        ["  II. Receivables", assets["C_receivables_current"], assets["C_receivables_previous"]],
        ["  III. Investments", assets["C_investments_current"], assets["C_investments_previous"]],
        ["  IV. Cash and Cash Equivalents", assets["C_cash_current"], assets["C_cash_previous"]],
        ["  Total for Section C", assets["C_total_current_assets_current"], assets["C_total_current_assets_previous"]],
        ["D. Prepaid Expenses", assets["D_prepaid_expenses_current"], assets["D_prepaid_expenses_previous"]],
        ["TOTAL ASSETS", assets["total_assets_current"], assets["total_assets_previous"]],
    ]
    for row in asset_data:
        cells = asset_table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = format_number(row[1])
        cells[2].text = format_number(row[2])

    # **Liabilities Section**
    doc.add_heading("Liabilities", level=2)
    liability_table = doc.add_table(rows=1, cols=3)
    liability_table.style = 'Table Grid'

    # Add headers
    headers = liability_table.rows[0].cells
    headers[0].text = "Sections, Groups, Items"
    headers[1].text = "Amount (Current Year)"
    headers[2].text = "Amount (Previous Year)"

    # Add liability data
    liability_data = [
        ["A. Equity", "", ""],
        ["  I. Issued Capital", liabilities["A_issued_capital_current"], liabilities["A_issued_capital_previous"]],
        ["  II. Share Premiums", liabilities["A_share_premiums_current"], liabilities["A_share_premiums_previous"]],
        ["  III. Revaluation Reserve", liabilities["A_revaluation_reserve_current"], liabilities["A_revaluation_reserve_previous"]],
        ["  IV. Reserves", liabilities["A_reserves_current"], liabilities["A_reserves_previous"]],
        ["  V. Retained Earnings (Loss)", liabilities["A_retained_earnings_current"], liabilities["A_retained_earnings_previous"]],
        ["  VI. Current Profit (Loss)", liabilities["A_current_profit_loss_current"], liabilities["A_current_profit_loss_previous"]],
        ["  VII. Shareholders Equity", liabilities["A_shareholders_equity_current"], liabilities["A_shareholders_equity_previous"]],
        ["  Total for Section A", liabilities["A_total_equity_current"], liabilities["A_total_equity_previous"]],
        ["B. Provisions and Similar Obligations", liabilities["B_provisions_current"], liabilities["B_provisions_previous"]],
        ["C. Liabilities", "", ""],
        ["  Up to 1 Year", liabilities["C_liabilities_one_year_current"], liabilities["C_liabilities_one_year_previous"]],
        ["  Over 1 Year", liabilities["C_liabilities_over_one_year_current"], liabilities["C_liabilities_over_one_year_previous"]],
        ["D. Financing and Deferred Income", liabilities["D_deferred_income_current"], liabilities["D_deferred_income_previous"]],
        ["E. Deferred Tax Liabilities", liabilities["E_deferred_tax_liabilities_current"], liabilities["E_deferred_tax_liabilities_previous"]],
        ["TOTAL LIABILITIES", liabilities["total_liabilities_current"], liabilities["total_liabilities_previous"]],
    ]
    for row in liability_data:
        cells = liability_table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = format_number(row[1])
        cells[2].text = format_number(row[2])

    # Save the document to a buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Return the response
    response = make_response(buffer.read())
    response.headers["Content-Disposition"] = "attachment; filename=balance_sheet.docx"
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return response



from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from flask import make_response
def export_balance_sheet_excel():
    # Fetch data using the calculation functions
    transactions = Transaction.query.filter_by(user_id=current_user.id).all()

    assets = calculate_asset_data(transactions)
    liabilities = calculate_liability_data(transactions)

    # Create a new Excel workbook and active sheet
    wb = Workbook()
    ws = wb.active
    ws.title = "Balance Sheet"

    # Title
    ws.merge_cells("A1:C1")
    title_cell = ws["A1"]
    title_cell.value = "Balance Sheet"
    title_cell.font = Font(size=16, bold=True)
    title_cell.alignment = Alignment(horizontal="center")

    # Custom formatting for numbers
    def format_number(value):
        if isinstance(value, (int, float)):
            return f"({abs(value):,.2f})" if value < 0 else f"{value:,.2f}"
        return value

    # **Assets Section**
    ws.append([""])  # Blank row for spacing
    ws.append(["Assets"])
    ws["A3"].font = Font(size=14, bold=True)

    # Add asset headers
    ws.append(["Sections, Groups, Items", "Amount (Current Year)", "Amount (Previous Year)"])
    header_row = ws[4]
    for cell in header_row:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")

    # Add asset data
    asset_data = [
        ["A. Subscribed but Unpaid Capital", assets["A_asset_unpaid_capital_current"], assets["A_asset_unpaid_capital_previous"]],
        ["B. Non-Current (Long-Term) Assets", "", ""],
        ["  I. Intangible Assets", assets["B_intangible_assets_current"], assets["B_intangible_assets_previous"]],
        ["  II. Property, Plant, and Equipment", assets["B_fixed_assets_current"], assets["B_fixed_assets_previous"]],
        ["  III. Long-Term Financial Assets", assets["B_long_term_financial_assets_current"], assets["B_long_term_financial_assets_previous"]],
        ["  IV. Deferred Taxes", assets["B_deferred_taxes_current"], assets["B_deferred_taxes_previous"]],
        ["  Total for Section B", assets["B_total_noncurrent_assets_current"], assets["B_total_noncurrent_assets_previous"]],
        ["C. Current (Short-Term) Assets", "", ""],
        ["  I. Inventory", assets["C_inventory_current"], assets["C_inventory_previous"]],
        ["  II. Receivables", assets["C_receivables_current"], assets["C_receivables_previous"]],
        ["  III. Investments", assets["C_investments_current"], assets["C_investments_previous"]],
        ["  IV. Cash and Cash Equivalents", assets["C_cash_current"], assets["C_cash_previous"]],
        ["  Total for Section C", assets["C_total_current_assets_current"], assets["C_total_current_assets_previous"]],
        ["D. Prepaid Expenses", assets["D_prepaid_expenses_current"], assets["D_prepaid_expenses_previous"]],
        ["TOTAL ASSETS", assets["total_assets_current"], assets["total_assets_previous"]],
    ]
    for row in asset_data:
        ws.append([row[0], format_number(row[1]), format_number(row[2])])

    # **Liabilities Section**
    ws.append([""])  # Blank row for spacing
    ws.append(["Liabilities"])
    ws[f"A{len(ws['A'])}"].font = Font(size=14, bold=True)

    # Add liability headers
    ws.append(["Sections, Groups, Items", "Amount (Current Year)", "Amount (Previous Year)"])
    header_row = ws[len(ws['A'])]
    for cell in header_row:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")

    # Add liability data
    liability_data = [
        ["A. Equity", "", ""],
        ["  I. Issued Capital", liabilities["A_issued_capital_current"], liabilities["A_issued_capital_previous"]],
        ["  II. Share Premiums", liabilities["A_share_premiums_current"], liabilities["A_share_premiums_previous"]],
        ["  III. Revaluation Reserve", liabilities["A_revaluation_reserve_current"], liabilities["A_revaluation_reserve_previous"]],
        ["  IV. Reserves", liabilities["A_reserves_current"], liabilities["A_reserves_previous"]],
        ["  V. Retained Earnings (Loss)", liabilities["A_retained_earnings_current"], liabilities["A_retained_earnings_previous"]],
        ["  VI. Current Profit (Loss)", liabilities["A_current_profit_loss_current"], liabilities["A_current_profit_loss_previous"]],
        ["  VII. Shareholders Equity", liabilities["A_shareholders_equity_current"], liabilities["A_shareholders_equity_previous"]],
        ["  Total for Section A", liabilities["A_total_equity_current"], liabilities["A_total_equity_previous"]],
        ["B. Provisions and Similar Obligations", liabilities["B_provisions_current"], liabilities["B_provisions_previous"]],
        ["C. Liabilities", "", ""],
        ["  Up to 1 Year", liabilities["C_liabilities_one_year_current"], liabilities["C_liabilities_one_year_previous"]],
        ["  Over 1 Year", liabilities["C_liabilities_over_one_year_current"], liabilities["C_liabilities_over_one_year_previous"]],
        ["D. Financing and Deferred Income", liabilities["D_deferred_income_current"], liabilities["D_deferred_income_previous"]],
        ["E. Deferred Tax Liabilities", liabilities["E_deferred_tax_liabilities_current"], liabilities["E_deferred_tax_liabilities_previous"]],
        ["TOTAL LIABILITIES", liabilities["total_liabilities_current"], liabilities["total_liabilities_previous"]],
    ]
    for row in liability_data:
        ws.append([row[0], format_number(row[1]), format_number(row[2])])

    # Save the workbook to a buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    # Return the response
    response = make_response(buffer.read())
    response.headers["Content-Disposition"] = "attachment; filename=balance_sheet.xlsx"
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return response




# # Export functions
# def export_to_excel(transactions):
#     output = BytesIO()
#     workbook = xlsxwriter.Workbook(output, {'in_memory': True})
#     worksheet = workbook.add_worksheet()

#     # Write headers
#     headers = ['Date', 'Type', 'Category', 'Amount', 'Description']
#     for col_num, header in enumerate(headers):
#         worksheet.write(0, col_num, header)

#     # Write transaction data
#     for row_num, transaction in enumerate(transactions, start=1):
#         worksheet.write(row_num, 0, transaction.date.strftime('%Y-%m-%d'))
#         worksheet.write(row_num, 1, transaction.type)
#         worksheet.write(row_num, 2, transaction.category)
#         worksheet.write(row_num, 3, transaction.amount)
#         worksheet.write(row_num, 4, transaction.description)

#     workbook.close()
#     output.seek(0)
#     return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
#                      as_attachment=True, download_name='Income_Expense_Report.xlsx')


# def export_to_pdf(transactions):
#     # Generate PDF using matplotlib or reportlab
#     output = BytesIO()
#     plt.figure(figsize=(8, 6))
#     categories = [t.category for t in transactions]
#     amounts = [t.amount for t in transactions]
#     plt.bar(categories, amounts)
#     plt.title('Expenses by Category')
#     plt.xlabel('Categories')
#     plt.ylabel('Amount')
#     plt.tight_layout()
#     plt.savefig(output, format='pdf')
#     output.seek(0)
#     return send_file(output, mimetype='application/pdf', as_attachment=True, download_name='Income_Expense_Report.pdf')

    

# def export_to_word(transactions):
#     document = Document()
#     document.add_heading('Income and Expense Report', level=1)
#     table = document.add_table(rows=1, cols=5)
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'Date'
#     hdr_cells[1].text = 'Type'
#     hdr_cells[2].text = 'Category'
#     hdr_cells[3].text = 'Amount'
#     hdr_cells[4].text = 'Description'

#     for transaction in transactions:
#         row_cells = table.add_row().cells
#         row_cells[0].text = transaction.date.strftime('%Y-%m-%d')
#         row_cells[1].text = transaction.type
#         row_cells[2].text = transaction.category
#         row_cells[3].text = str(transaction.amount)
#         row_cells[4].text = transaction.description

#     output = BytesIO()
#     document.save(output)
#     output.seek(0)
#     return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
#                      as_attachment=True, download_name='Income_Expense_Report.docx')
